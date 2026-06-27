"""export.py — turn the tailored Markdown resume into ATS-friendly DOCX / PDF.

Both outputs are deliberately plain: single column, standard fonts, selectable
text, no tables/images/graphics — which is exactly what ATS parsers read best.
"""

import io
import re


def _segments(text):
    """Split a line into (text, is_bold) segments on **bold** markers."""
    out = []
    for p in re.split(r"(\*\*.+?\*\*)", text):
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            out.append((p[2:-2], True))
        else:
            out.append((p, False))
    return out


def md_to_docx_bytes(md):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for raw in md.splitlines():
        s = raw.strip()
        if not s:
            continue
        if s.startswith("> "):
            s = s[2:].strip()
        if s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=0)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=1)
        elif s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=2)
        elif s.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            for t, b in _segments(s[2:].strip()):
                p.add_run(t).bold = b
        else:
            p = doc.add_paragraph()
            for t, b in _segments(s):
                p.add_run(t).bold = b

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def md_to_pdf_bytes(md):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from xml.sax.saxutils import escape

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, title="Resume",
                            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
                            topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    base = getSampleStyleSheet()["Normal"]
    body = ParagraphStyle("body", parent=base, fontName="Helvetica",
                          fontSize=9.5, leading=13, spaceAfter=3)
    h0 = ParagraphStyle("h0", parent=body, fontName="Helvetica-Bold", fontSize=16, spaceAfter=2)
    h1 = ParagraphStyle("h1", parent=body, fontName="Helvetica-Bold", fontSize=11.5,
                        spaceBefore=9, spaceAfter=3)
    h2 = ParagraphStyle("h2", parent=body, fontName="Helvetica-Bold", fontSize=10,
                        spaceBefore=5, spaceAfter=1)
    bul = ParagraphStyle("bul", parent=body, leftIndent=14, bulletIndent=2, spaceAfter=2)

    def fmt(t):
        return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escape(t))

    flow = []
    for raw in md.splitlines():
        s = raw.strip()
        if not s:
            flow.append(Spacer(1, 4))
            continue
        if s.startswith("> "):
            s = s[2:].strip()
        if s.startswith("# "):
            flow.append(Paragraph(fmt(s[2:].strip()), h0))
        elif s.startswith("## "):
            flow.append(Paragraph(fmt(s[3:].strip()), h1))
        elif s.startswith("### "):
            flow.append(Paragraph(fmt(s[4:].strip()), h2))
        elif s.startswith(("- ", "* ", "• ")):
            flow.append(Paragraph(fmt(s[2:].strip()), bul, bulletText="•"))
        else:
            flow.append(Paragraph(fmt(s), body))

    doc.build(flow)
    return buf.getvalue()
